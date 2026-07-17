from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "competition-documents"
SRC = ROOT / "docs" / "competition"
OUT.mkdir(parents=True, exist_ok=True)
SRC.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK = "1F4D78"
NAVY = "16324F"
MUTED = "687386"
LIGHT = "F2F4F7"
PALE = "E8EEF5"
GREEN = "18866B"
GOLD = "8A6500"
WHITE = "FFFFFF"
TOTAL_DXA = 9360


def set_run_font(run, size=11, bold=False, color="20242A", italic=False, name="Noto Sans CJK SC"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_doc(title):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(.492)
    section.footer_distance = Inches(.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Noto Sans CJK SC"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans CJK SC")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans CJK SC")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Noto Sans CJK SC"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans CJK SC")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans CJK SC")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Noto Sans CJK SC"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans CJK SC")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans CJK SC")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(.5)
        style.paragraph_format.first_line_indent = Inches(-.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run(f"问阶 AI 学习工作台  |  {title}")
    set_run_font(hr, 9, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("华中师范大学 · 张书旋 · 2023214382")
    set_run_font(fr, 9, color=MUTED)
    return doc


def add_cover(doc, title, subtitle, doc_no):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("问阶 AI 学习工作台")
    set_run_font(r, 13, bold=True, color=GREEN)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, 28, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, 13, color=MUTED)
    p.paragraph_format.space_after = Pt(46)
    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [2400, 6960])
    values = [
        ("文档编号", doc_no),
        ("版本", "V1.0"),
        ("编制单位", "华中师范大学"),
        ("编制日期", "2026年7月17日"),
    ]
    for row, (label, value) in zip(table.rows, values):
        set_cell_shading(row.cells[0], PALE)
        for idx, text in enumerate((label, value)):
            p = row.cells[idx].paragraphs[0]
            rr = p.add_run(text)
            set_run_font(rr, 10.5, bold=(idx == 0), color=NAVY if idx == 0 else "20242A")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("内部参赛提交稿")
    set_run_font(r, 10, bold=True, color=GOLD)
    doc.add_page_break()


def add_control(doc, purpose):
    doc.add_heading("文档控制", 1)
    add_table(doc, ["项目", "内容"], [
        ["文档目的", purpose],
        ["适用范围", "参赛提交、开发交接、系统验收和答辩说明"],
        ["事实基线", "项目源代码、README、接口实现、知识库OCR结果及本地测试记录"],
        ["数据说明", "未开展大规模正式问卷；调研报告不虚构样本量和统计结论"],
        ["保密与版权", "用户上传PDF仅存于本地私有目录，不随仓库分发"],
    ], [2200, 7160])
    doc.add_heading("修订记录", 2)
    add_table(doc, ["版本", "日期", "修订内容", "编制人"], [
        ["V1.0", "2026-07-17", "首次形成参赛标准文档", "张书旋"],
    ], [1000, 1700, 5260, 1400])


def add_para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r)


def add_table(doc, headers, rows, widths=None):
    widths = widths or [TOTAL_DXA // len(headers)] * len(headers)
    widths[-1] += TOTAL_DXA - sum(widths)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_header(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], PALE)
        p = table.rows[0].cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, 9.5, bold=True, color=NAVY)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, 9.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_flow(doc, title, stages):
    # This is the first subsection below a chapter heading; keep the semantic
    # outline contiguous for screen readers even though it is visually compact.
    doc.add_heading(title, 2)
    table = doc.add_table(rows=1, cols=len(stages))
    widths = [TOTAL_DXA // len(stages)] * len(stages)
    widths[-1] += TOTAL_DXA - sum(widths)
    set_table_geometry(table, widths)
    for idx, stage in enumerate(stages):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, PALE if idx % 2 == 0 else "F7FAFC")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(stage + ("  →" if idx < len(stages) - 1 else ""))
        set_run_font(r, 9, bold=True, color=DARK)
    doc.add_paragraph()


def add_callout(doc, label, text, color=GREEN):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TOTAL_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F8F7")
    p = cell.paragraphs[0]
    r = p.add_run(label + "：")
    set_run_font(r, 10, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, 10)
    doc.add_paragraph()


def save_markdown(name, title, sections):
    lines = [f"# {title}", "", "> 问阶 AI 学习工作台 · 参赛提交稿", ""]
    for heading, body in sections:
        lines.append(f"## {heading}")
        lines.append("")
        if isinstance(body, list):
            lines.extend([f"- {item}" for item in body])
        else:
            lines.append(body)
        lines.append("")
    (SRC / name).write_text("\n".join(lines), encoding="utf-8")


def build_development():
    title = "系统开发说明书"
    doc = configure_doc(title)
    add_cover(doc, title, "多智能体个性化高校学习资源系统", "WJ-AI-DEV-001")
    add_control(doc, "系统说明从需求、架构、模块、数据、接口、安全、部署和运维层面描述可复现的开发全过程。")
    doc.add_heading("1 项目概述", 1)
    add_para(doc, "问阶 AI 学习工作台面向高校专业课程自主学习场景，解决资源繁杂、难以匹配个体基础、教师集体讲授难以兼顾学习节奏等问题。系统以动态学习画像为入口，通过真实独立调用的多智能体链路，结合本地课程PDF知识库与可信高校资源目录，形成资源生成、路径规划、练习反馈、学习评估和动态优化闭环。")
    add_callout(doc, "交付状态", "核心功能、非功能机制、课程知识库、测试数据和提交声明均已落地；AI生成内容仍需回到教材、教师或原始文献复核。")

    doc.add_heading("2 需求分析", 1)
    doc.add_heading("2.1 核心用户问题", 2)
    add_bullets(doc, [
        "课程资料分散在教材、课程平台、论文、视频和个人笔记中，检索成本高。",
        "同一课程内学生的知识基础、目标、节奏、认知风格和易错点不同。",
        "大模型可生成内容，但存在来源不明、事实错误和不同课程内容串扰风险。",
        "资源生成后缺乏可执行路径、掌握证据和根据学习效果持续调整的机制。",
    ])
    doc.add_heading("2.2 功能需求映射", 2)
    add_table(doc, ["需求", "系统能力", "验收证据"], [
        ["画像构建", "自然对话抽取8维画像并持续更新", "画像页面与持久化数据"],
        ["多智能体资源", "分析、检索、文档、导图、命题、阅读、代码、PPT、审核", "独立调用轨迹与资源卡"],
        ["课程知识库", "上传PDF、文本解析/OCR、切分、检索、页码引用", "441页教材、494片段"],
        ["学习路径", "4至6阶段、Todo、顺序理由、掌握证据", "路径页和完成状态"],
        ["效果评估", "练习、错题、资源使用、路径执行等多维评估", "复盘页与动态策略"],
        ["安全合规", "前后端过滤、引用风险检测、来源区分", "CONTENT_SAFETY_BLOCKED"],
    ], [1800, 4560, 3000])

    doc.add_heading("3 总体架构", 1)
    add_flow(doc, "3.1 系统主链路", ["用户交互", "画像与需求", "课程RAG", "多Agent生成", "路径推送", "评估优化"])
    add_para(doc, "系统采用浏览器前端、Node.js服务层、MySQL用户数据层、本地课程知识库与外部模型服务组成的分层架构。前端负责交互、流式呈现和本地体验；服务端负责认证、数据接口、模型代理、内容安全、PDF知识库和PPT生成；MySQL保存账户和学习数据；`.runtime` 保存不入库、不提交Git的私有PDF与OCR索引。")
    add_table(doc, ["层级", "主要技术", "职责"], [
        ["表现层", "HTML/CSS/JavaScript", "对话、画像、资源、路径、复盘、移动端"],
        ["应用层", "Node.js HTTP服务", "认证、路由、代理、知识库、文件生成"],
        ["智能层", "独立Agent模型调用", "分析、生成、审核、路径和评估"],
        ["数据层", "MySQL + 浏览器缓存", "用户、会话、学习状态和行为证据"],
        ["知识层", "pdf-parse/Poppler/Tesseract", "课程PDF解析、OCR、切分和本地检索"],
    ], [1500, 2800, 5060])

    doc.add_heading("4 多智能体设计", 1)
    add_flow(doc, "4.1 真实协作流程", ["需求分析Agent", "课程RAG", "资源Agent并行", "审核整合Agent", "动态学习路径"])
    add_table(doc, ["智能体", "输入", "输出", "失败策略"], [
        ["需求分析师", "需求、画像、学习信号、课程片段", "主题、目标、先修与短板", "45秒超时后重试"],
        ["教育资源检索", "课程主题", "可信高校平台入口", "本地零成本降级"],
        ["文档/导图/题库/阅读/代码/PPT", "分析结果、画像、RAG片段", "独立资源中间产物", "并行、失败跳过"],
        ["审核整合", "所有实际中间结果", "审核意见、路径和反馈", "错误显式呈现"],
        ["效果评估", "行为、正确率、错题、使用和路径", "多维评分与调整策略", "规则评估兜底"],
    ], [1700, 2500, 3300, 1860])
    add_para(doc, "区别于单次提示词中的角色模拟，系统对需求分析、各资源类型和最终审核分别发起独立模型请求，并保存角色、开始与完成时间、尝试次数、状态、中间结果摘要和反馈。资源Agent使用 Promise.allSettled 并行，单个失败不会中断整轮。")

    doc.add_heading("5 课程知识库与RAG", 1)
    add_flow(doc, "5.1 PDF处理流程", ["上传PDF", "检测文本层", "扫描版OCR", "按页切分", "相关度检索", "注入Agent并引用"])
    add_table(doc, ["指标", "当前结果"], [
        ["课程", "机器学习（高校计算机/人工智能专业核心课程）"],
        ["初始文档", "周志华《机器学习》本地私有PDF"],
        ["页数", "441页"],
        ["解析方式", "扫描版中文OCR"],
        ["检索片段", "494个"],
        ["课程结构", "16章、64建议学时、6个建议实验、综合项目考核"],
        ["测试结果", "12/12检索测试通过"],
    ], [2500, 6860])
    add_para(doc, "知识库片段保留课程、文档、页码、章节位置和提取方式。对话与资源Agent必须优先使用召回片段，并以“课程文档：文档名，PDF第N页”标注；未覆盖内容标记为“AI扩展”。系统只向模型发送与当前问题相关的少量片段，不发送整本电子书。")

    doc.add_heading("6 功能模块", 1)
    add_table(doc, ["模块", "关键功能"], [
        ["账户与管理", "注册登录、HttpOnly会话、管理员、公告、审计"],
        ["对话辅导", "流式输出、多轮上下文、图片输入、Markdown、代码高亮"],
        ["学生画像", "8维画像、确定/推测/待补充、活动触发更新"],
        ["资源中心", "7类资源、Agent选择、题量配置、来源和安全标识"],
        ["路径与推送", "分知识大类路径、Todo、阶段推荐、优先级调整"],
        ["复盘评估", "练习正确率、错题、资源使用、路径完成度、多维评估"],
        ["知识库", "PDF上传、OCR、目录、检索、页码引用"],
    ], [2100, 7260])

    doc.add_heading("7 数据与接口", 1)
    doc.add_heading("7.1 主要数据实体", 2)
    add_bullets(doc, [
        "wj_users：用户、角色、状态与最近登录。",
        "wj_sessions：服务端会话与有效期。",
        "wj_user_data：画像、资源、路径、错题、对话和行为等用户数据。",
        ".runtime/knowledge-base：私有PDF、meta.json和chunks.json。",
    ])
    doc.add_heading("7.2 关键接口", 2)
    add_table(doc, ["方法", "接口", "用途"], [
        ["POST", "/api/chat", "模型代理、流式/非流式对话和服务端安全检查"],
        ["POST", "/api/knowledge-base/upload", "上传最大80MB课程PDF"],
        ["GET", "/api/knowledge-base", "知识库目录"],
        ["GET", "/api/knowledge-base/search", "本地课程片段检索"],
        ["GET", "/api/knowledge-base/tasks/:id", "OCR任务进度"],
        ["POST", "/api/presentations", "生成PPT文件或异步任务"],
        ["GET/PUT", "/api/user-data", "读取和保存学习数据"],
    ], [1000, 3700, 4660])

    doc.add_heading("8 安全、性能与可靠性", 1)
    add_bullets(doc, [
        "输入安全：浏览器和服务端双层检查现实伤害、凭证盗取和恶意软件等操作性请求。",
        "Prompt注入：检测覆盖系统规则或索取隐藏提示的请求，保持系统规则优先。",
        "输出安全：Agent结果进入资源池前检查；论文、URL和绝对化结论标记待复核。",
        "响应控制：对话流式输出；资源Agent并行；单次45秒超时、自动重试一次、失败降级。",
        "进度体验：显示整体百分比、逐Agent执行状态和OCR页数。",
        "版权与隐私：课程PDF在`.runtime`本地私有保存，不随仓库分发。",
    ])

    doc.add_heading("9 部署与运行", 1)
    add_steps(doc, [
        "安装Node.js依赖并确认Poppler、Tesseract及chi_sim语言包可用。",
        "创建MySQL数据库并配置DATABASE_URL或MYSQL_*环境变量。",
        "配置模型代理密钥、基础地址和模型名称；视频能力默认不启用。",
        "运行 npm run dev，访问本地服务。",
        "在资源页上传课程PDF，等待OCR或文本解析完成。",
        "运行 npm run test:knowledge-base 验证课程检索。",
    ])
    doc.add_heading("10 创新点与边界", 1)
    add_bullets(doc, [
        "自然对话驱动的8维动态画像，避免一次性繁琐问卷。",
        "独立调用、并行执行、可追踪和可审核的真实多智能体链路。",
        "扫描版中文教材OCR、页码级RAG和AI扩展标记。",
        "从资源生成到路径、练习、错题、评估和再推送的闭环。",
        "检索资源、课程知识库和AI生成资源三类证据明确区分。",
    ])
    add_callout(doc, "已知边界", "规则安全审核不能替代生产级审核服务；OCR存在识别误差；大模型输出仍需人工和原始来源复核；正式用户调研需补充真实样本。", GOLD)
    return doc


def build_testing():
    title = "系统测试说明书"
    doc = configure_doc(title)
    add_cover(doc, title, "功能、接口、知识库、安全、性能与移动端验证", "WJ-AI-TEST-001")
    add_control(doc, "定义测试范围、环境、用例、执行证据、结果和遗留风险，支持项目验收与答辩。")
    doc.add_heading("1 测试目标与范围", 1)
    add_para(doc, "测试验证核心业务闭环、真实多智能体协作、课程PDF知识库、内容安全、响应控制和移动端适配。测试以本地代码、接口返回、浏览器控制台、OCR结果和自动化脚本为证据，不把模型主观质量等同于确定性软件通过。")
    add_table(doc, ["测试域", "覆盖内容", "结论"], [
        ["功能", "对话、画像、资源、路径、推送、复盘、账户", "核心功能具备"],
        ["知识库", "上传、OCR、切分、检索、页码和无关查询", "12/12通过"],
        ["安全", "前端拦截、服务端拦截、输出审计", "危险请求已拦截"],
        ["性能", "流式、并行、超时、重试、降级、进度", "机制已实现"],
        ["兼容", "桌面与390×844移动视口", "无横向溢出"],
    ], [1500, 5700, 2160])

    doc.add_heading("2 测试环境", 1)
    add_table(doc, ["项目", "配置"], [
        ["操作系统", "macOS，本地开发环境"],
        ["运行时", "Node.js 24.x；浏览器原生JavaScript"],
        ["数据库", "MySQL；前端缓存用于体验优化"],
        ["PDF工具", "pdf-parse 1.1.1、Poppler pdftoppm、Tesseract 5.5.2、chi_sim"],
        ["测试课程", "机器学习；441页扫描PDF；494检索片段"],
        ["移动视口", "390×844"],
        ["测试日期", "2026-07-17"],
    ], [2200, 7160])

    doc.add_heading("3 测试策略", 1)
    add_steps(doc, [
        "静态检查：JavaScript语法、Git差异格式和依赖完整性。",
        "接口测试：知识库目录、上传、任务进度、检索和内容安全返回码。",
        "功能测试：按用户主流程验证页面和状态持久化。",
        "数据驱动测试：使用12条机器学习检索用例运行自动脚本。",
        "响应测试：验证流式、并行、进度、超时、重试和降级机制。",
        "移动端回归：在390×844视口检查布局和横向溢出。",
    ])

    doc.add_heading("4 功能测试用例", 1)
    functional_cases = [
        ["F-01", "自然对话", "发送课程问题", "流式回答并保存历史", "通过"],
        ["F-02", "学习画像", "完成多轮学习对话", "8维画像随学更新", "通过"],
        ["F-03", "资源Agent", "选择文档、题库、代码", "独立并行生成并记录轨迹", "通过"],
        ["F-04", "资源失败", "模拟单Agent失败", "其他Agent继续，失败显式显示", "通过"],
        ["F-05", "学习路径", "生成同一知识大类资源", "保留有效阶段并增量更新", "通过"],
        ["F-06", "错题复盘", "标记错误练习", "进入错题与评估证据", "通过"],
        ["F-07", "资源推送", "打开推送页", "按阶段和评估排序", "通过"],
        ["F-08", "PPT", "选择PPT Agent", "生成任务或下载文件", "通过/依配置"],
        ["F-09", "主题", "切换深浅主题并刷新", "主题保持", "通过"],
        ["F-10", "移动端", "390×844打开资源页", "单列布局且无横向溢出", "通过"],
    ]
    add_table(doc, ["编号", "模块", "操作", "预期", "结果"], functional_cases, [900, 1500, 2450, 3210, 1300])

    doc.add_heading("5 知识库测试", 1)
    add_para(doc, "测试数据来自 `knowledge-base-tests/machine-learning-evaluation.json`，自动脚本为 `scripts/evaluate-knowledge-base.js`。通过条件包括相关问题有召回、无关问题不误召回、结果携带PDF页码。")
    kb_cases = [
        ["K-01", "经验误差与泛化误差", "5", "283", "通过"],
        ["K-02", "十折交叉验证", "5", "42", "通过"],
        ["K-03", "线性回归与最小二乘", "5", "70", "通过"],
        ["K-04", "决策树信息增益", "5", "103", "通过"],
        ["K-05", "BP神经网络", "5", "120", "通过"],
        ["K-06", "支持向量机核函数", "5", "144", "通过"],
        ["K-07", "朴素贝叶斯", "5", "170", "通过"],
        ["K-08", "Bagging与Boosting", "5", "13", "通过"],
        ["K-09", "K均值聚类", "5", "218", "通过"],
        ["K-10", "主成分分析", "5", "245", "通过"],
        ["K-11", "强化学习", "5", "389", "通过"],
        ["K-12", "唐代诗歌平仄（无关）", "0", "-", "通过"],
    ]
    add_table(doc, ["编号", "查询主题", "召回数", "首条页码", "结果"], kb_cases, [900, 3600, 1400, 1800, 1660])
    add_callout(doc, "实际结果", "知识库检索测试12/12通过。OCR共处理441页，生成494个检索片段。")

    doc.add_heading("6 内容安全测试", 1)
    safety_cases = [
        ["S-01", "危险软件与凭证盗取组合请求", "前端保持输入、不发起生成", "通过"],
        ["S-02", "同一危险请求直接调用/api/chat", "HTTP 400 + CONTENT_SAFETY_BLOCKED", "通过"],
        ["S-03", "要求泄露系统提示词", "增加系统防注入约束", "通过"],
        ["S-04", "AI资源含外部URL或论文线索", "标记引用待复核", "通过"],
        ["S-05", "检索资源与AI生成资源", "使用不同来源徽标", "通过"],
    ]
    add_table(doc, ["编号", "输入/场景", "预期", "结果"], safety_cases, [900, 3600, 3560, 1300])

    doc.add_heading("7 性能与可靠性验证", 1)
    add_table(doc, ["机制", "设计阈值/策略", "验证"], [
        ["对话响应", "SSE流式读取与逐字呈现", "已实现"],
        ["资源并发", "资源Agent使用Promise.allSettled", "已实现"],
        ["单Agent超时", "45秒", "代码检查通过"],
        ["重试", "失败自动重试1次", "代码检查通过"],
        ["降级", "重试失败后跳过，不中断其他资源", "代码检查通过"],
        ["进度", "总体百分比、逐Agent状态、OCR页数", "界面已实现"],
        ["PDF大小", "最大80MB", "服务端限制已实现"],
    ], [1900, 4300, 3160])
    add_para(doc, "说明：模型端到端耗时受服务商、网络和输出长度影响。本轮未宣称固定平均响应毫秒数；正式答辩环境应补充连续多轮计时数据。")

    doc.add_heading("8 缺陷、风险与回归建议", 1)
    add_table(doc, ["风险", "影响", "当前措施", "后续建议"], [
        ["OCR错字", "检索或引用偏差", "保留页码与原PDF", "增加版面OCR和人工校验"],
        ["模型幻觉", "学术事实错误", "RAG、来源标识、审核Agent", "增加DOI/Crossref验证"],
        ["模型服务波动", "生成超时", "重试和降级", "增加熔断、队列和缓存"],
        ["规则安全覆盖有限", "隐晦风险漏检", "前后端双层规则", "接入生产级审核API"],
        ["正式调研样本不足", "需求证据不充分", "场景研究透明披露", "开展问卷与访谈"],
    ], [1900, 1900, 2700, 2860])
    doc.add_heading("9 验收结论", 1)
    add_para(doc, "在当前本地测试环境中，系统核心功能、课程知识库、安全拦截、移动端布局和多智能体可靠性机制达到参赛演示条件。知识库自动测试12/12通过。正式上线前仍需完成更大规模用户调研、持续性能压测、OCR抽样人工核验和生产级安全审核接入。")
    return doc


def build_research():
    title = "用户需求调研与分析报告"
    doc = configure_doc(title)
    add_cover(doc, title, "新时代大学生专业课程学习需求与AI技术映射", "WJ-AI-REQ-001")
    add_control(doc, "说明目标用户、学习场景、核心痛点、需求优先级和技术结合点，为系统设计提供可追溯依据。")
    doc.add_heading("1 研究说明", 1)
    add_callout(doc, "证据边界", "本报告基于竞赛任务要求、项目开发过程中的典型学习场景、系统行为证据和公开可观察的高校学习流程形成。尚未开展经伦理审批的大样本问卷或正式访谈，因此不虚构样本量、百分比或显著性结论。", GOLD)
    add_para(doc, "研究目标是回答三个问题：学生在专业课程自主学习中卡在哪里；哪些问题适合由AI和多智能体解决；如何避免生成式AI带来来源、事实、安全和依赖风险。")

    doc.add_heading("2 目标用户与场景", 1)
    add_table(doc, ["用户原型", "主要目标", "典型困难", "需要的支持"], [
        ["基础补齐型", "理解课程概念并完成作业", "先修不足、术语密集", "分层讲解、图解、基础题"],
        ["考试复习型", "短期构建知识框架", "资料分散、时间有限", "路径、重点、错题和测验"],
        ["项目实践型", "把知识迁移到代码和项目", "会概念不会实现", "代码骨架、测试和调试"],
        ["研究入门型", "阅读论文并定位方向", "文献门槛和来源可信度", "课程基础、检索入口和引用核验"],
        ["教师辅助型", "准备分层材料和练习", "难以兼顾不同学生", "资源组合、题量配置和效果证据"],
    ], [1800, 2200, 2500, 2860])
    doc.add_heading("2.1 核心使用旅程", 2)
    add_flow(doc, "学生学习旅程", ["提出问题", "系统理解画像", "检索课程教材", "生成个性化资源", "执行学习路径", "练习与复盘"])

    doc.add_heading("3 主要痛点分析", 1)
    pain_points = [
        ["P1", "资源繁杂无序", "课程平台、教材、短视频、论文和工具彼此割裂", "高"],
        ["P2", "个体匹配不足", "内容难度、例子类型和学习节奏无法自动适配", "高"],
        ["P3", "生成内容不可信", "链接、论文和确定性结论可能缺少证据", "高"],
        ["P4", "缺少行动顺序", "获得很多资料但不知道先学什么、如何判断掌握", "高"],
        ["P5", "反馈闭环薄弱", "错题、资源使用和路径完成没有共同驱动下一步", "高"],
        ["P6", "长任务等待", "PPT、题库和多类资源生成可能白屏等待", "中"],
        ["P7", "移动使用受限", "手机端卡片、长链接和按钮容易拥挤", "中"],
    ]
    add_table(doc, ["编号", "痛点", "表现", "优先级"], pain_points, [900, 2100, 5060, 1300])

    doc.add_heading("4 需求分层", 1)
    doc.add_heading("4.1 必须需求", 2)
    add_bullets(doc, [
        "通过自然语言持续构建不少于6维的动态学习画像。",
        "至少5类个性化资源，并明确多智能体协作过程。",
        "结合专业、进度、掌握情况和偏好规划动态路径。",
        "至少一门完整高校专业课程知识库作为初始输入。",
        "内容安全、来源区分、响应进度和失败恢复。",
    ])
    doc.add_heading("4.2 期望需求", 2)
    add_bullets(doc, [
        "即时文字、图片、代码和图解辅导。",
        "学习效果多维评估与资源策略动态调整。",
        "PPT、Word、Markdown和JSON等可交付文件。",
        "手机端可用、管理员可运营、个人数据可持久化。",
    ])

    doc.add_heading("5 技术与需求结合点", 1)
    mapping = [
        ["资源繁杂", "本地课程RAG + 可信高校平台目录", "召回教材页码并补充可溯源入口"],
        ["个体差异", "8维动态画像 + 需求轨迹", "调节深度、难度、风格和题目梯度"],
        ["材料单一", "独立资源Agent并行", "文档、导图、题库、阅读、代码和PPT"],
        ["路径缺失", "审核整合Agent + 动态路径", "4至6阶段、Todo和掌握证据"],
        ["反馈断裂", "行为、错题、正确率、资源使用和路径证据", "评估后调整推送与计划"],
        ["内容不可信", "来源徽标、页码引用、输出审核", "检索/教材/AI扩展明确区分"],
        ["等待焦虑", "流式、并行、实时进度、超时重试", "降低白屏和整轮失败"],
    ]
    add_table(doc, ["用户需求", "技术方案", "结合方式"], mapping, [1900, 3200, 4260])

    doc.add_heading("6 典型场景案例", 1)
    doc.add_heading("6.1 机器学习课程补弱", 2)
    add_para(doc, "学生提出“支持向量机为什么需要核函数”。系统从本地《机器学习》PDF召回相关页码，结合画像确定讲解深度；文档Agent给出概念与公式解释，导图Agent组织线性可分、间隔、核映射关系，题库Agent生成分层练习，代码Agent提供可运行实验；审核Agent把资源安排到诊断、理解、练习、实操和复盘阶段。")
    doc.add_heading("6.2 考试前快速复习", 2)
    add_para(doc, "学生说明考试日期、基础和薄弱章节。系统记录目标与时间约束，从课程知识库召回章节内容，生成重点文档、导图和题库，路径根据错题与正确率动态调整。已完成资源降权，未掌握知识点优先推送。")
    doc.add_heading("6.3 无关或高风险问题", 2)
    add_para(doc, "无关问题不应强行召回机器学习教材；操作性危险请求在浏览器和服务端被拦截。系统对论文、外部链接和绝对化学术结论显示复核提示，避免把AI输出当作确定事实。")

    doc.add_heading("7 需求优先级与验收指标", 1)
    add_table(doc, ["优先级", "需求", "验收指标"], [
        ["P0", "动态画像", "不少于6维；当前实现8维"],
        ["P0", "多智能体资源", "至少5类；当前提供7类并记录独立轨迹"],
        ["P0", "课程知识库", "至少1门完整课程；当前441页、494片段"],
        ["P0", "路径与推送", "阶段、顺序、Todo、掌握证据和推荐理由"],
        ["P0", "内容安全", "危险请求前后端拦截；来源明确区分"],
        ["P1", "评估闭环", "练习、错题、资源和路径证据进入评估"],
        ["P1", "移动体验", "390×844无横向溢出，主要按钮可触控"],
        ["P1", "响应控制", "流式、并行、进度、45秒超时和重试"],
    ], [1200, 3000, 5160])

    doc.add_heading("8 后续正式调研方案", 1)
    add_para(doc, "为将场景分析升级为可统计的用户研究，建议在提交前或赛后执行以下计划。")
    add_steps(doc, [
        "招募不同专业、年级和学习目标的学生，明确知情同意和匿名规则。",
        "使用半结构化访谈了解资源获取、AI使用、错题复盘和等待体验。",
        "设计5级量表问卷，测量资源匹配、可信度、认知负担和使用意愿。",
        "开展任务测试：完成教材问答、生成资源、执行路径和错题复盘。",
        "记录任务完成率、用时、错误、主观评分和开放反馈。",
        "比较使用系统前后的资源筛选时间和知识掌握证据，不夸大因果。",
    ])
    add_table(doc, ["建议指标", "定义"], [
        ["任务完成率", "在规定时间内完成指定学习任务的参与者比例"],
        ["资源匹配评分", "学生对难度、形式和当前目标匹配程度的5级评分"],
        ["来源信任度", "学生能否识别教材、检索资源和AI扩展"],
        ["路径可执行性", "Todo是否清晰、顺序是否合理、掌握证据是否可验证"],
        ["等待体验", "首个反馈时间、总耗时和进度信息满意度"],
    ], [2400, 6960])

    doc.add_heading("9 研究结论", 1)
    add_para(doc, "需求分析表明，项目价值不在单次生成更多资料，而在于把学生画像、课程证据、多智能体资源、可执行路径和学习反馈连接成持续优化系统。当前实现与核心需求高度一致，并通过真实课程PDF、检索测试和安全拦截提供了可核验证据。正式提交时应继续保持证据边界透明，并补充真实用户样本。")
    return doc


def main():
    docs = [
        ("问阶AI学习工作台_系统开发说明书.docx", build_development()),
        ("问阶AI学习工作台_系统测试说明书.docx", build_testing()),
        ("问阶AI学习工作台_用户需求调研与分析报告.docx", build_research()),
    ]
    for filename, doc in docs:
        doc.save(OUT / filename)
    save_markdown("系统开发说明书.md", "系统开发说明书", [
        ("项目定位", "面向高校专业课程学习的多智能体个性化学习工作台。"),
        ("架构摘要", "浏览器前端、Node.js服务、MySQL、本地课程RAG与外部模型服务。"),
        ("核心证据", ["8维画像", "真实独立Agent链路", "441页课程PDF与494片段", "12/12知识库检索测试", "内容安全与进度机制"]),
        ("完整内容", "请以同目录Word正式版为准。"),
    ])
    save_markdown("系统测试说明书.md", "系统测试说明书", [
        ("测试范围", "功能、接口、知识库、安全、性能和移动端。"),
        ("核心结果", ["知识库12/12通过", "危险请求服务端返回CONTENT_SAFETY_BLOCKED", "390×844无横向溢出", "JavaScript语法和Git差异格式检查通过"]),
        ("完整内容", "请以同目录Word正式版为准。"),
    ])
    save_markdown("用户需求调研与分析报告.md", "用户需求调研与分析报告", [
        ("证据边界", "未虚构大样本问卷；基于竞赛要求、典型学习场景、系统行为和实际测试证据。"),
        ("核心痛点", ["资源繁杂", "个体匹配不足", "内容可信度", "行动顺序缺失", "反馈闭环薄弱"]),
        ("完整内容", "请以同目录Word正式版为准。"),
    ])
    print("\n".join(str(OUT / name) for name, _ in docs))


if __name__ == "__main__":
    main()
